"""
RAG workflow services for processing content and generating embeddings.
"""

import os
import re
import json
import asyncio
import aiohttp
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime, timedelta
import numpy as np
from urllib.parse import urlparse, parse_qs
import tiktoken
import openai
from bs4 import BeautifulSoup
import PyPDF2
import docx
import csv
import io

from sqlalchemy.orm import Session
from sqlalchemy import update

from .models import (
    RAGWorkflow, RAGSource, RAGChunk, RAGEmbedding, RAGProcessingStep, 
    RAGWorkflowEvent, WorkflowStatus, SourceStatus, SourceType,
    VectorStore, EmbeddingModel
)
from ..core.logging import logger
from .external_search_service import ExternalSearchService
from .jina_service import JinaService


class RAGProcessingService:
    """Main service for processing RAG workflows."""
    
    def __init__(self, db: Session):
        self.db = db
        self.openai_client = openai.AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        self.embedding_cache = {}
        self.external_search_service = ExternalSearchService()
        self.jina_service = JinaService()
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
        
    async def start_workflow(self, workflow_id: str) -> RAGWorkflow:
        """Start processing a RAG workflow."""
        workflow = self.db.query(RAGWorkflow).filter_by(id=workflow_id).first()
        if not workflow:
            raise ValueError(f"Workflow {workflow_id} not found")
            
        if workflow.status not in [WorkflowStatus.QUEUED, WorkflowStatus.FAILED]:
            raise ValueError(f"Workflow is not in a startable state: {workflow.status}")
            
        # Update workflow status
        workflow.status = WorkflowStatus.SCRAPING
        workflow.started_at = datetime.utcnow()
        workflow.progress = 0.0
        self.db.commit()
        
        # Log event
        self._log_event(workflow_id, "started", {"previous_status": WorkflowStatus.QUEUED})
        
        # Start async processing
        asyncio.create_task(self._process_workflow(workflow_id))
        
        return workflow
        
    async def _process_workflow(self, workflow_id: str):
        """Main workflow processing pipeline."""
        try:
            # Step 1: Scraping
            await self._process_scraping(workflow_id)
            
            # Step 2: Chunking (part of embedding)
            await self._process_chunking(workflow_id)
            
            # Step 3: Embedding
            await self._process_embedding(workflow_id)
            
            # Step 4: Indexing
            await self._process_indexing(workflow_id)
            
            # Step 5: Validation
            await self._process_validation(workflow_id)
            
            # Mark as completed
            workflow = self.db.query(RAGWorkflow).filter_by(id=workflow_id).first()
            workflow.status = WorkflowStatus.COMPLETED
            workflow.progress = 100.0
            workflow.completed_at = datetime.utcnow()
            self.db.commit()
            
            self._log_event(workflow_id, "completed", {"duration_seconds": (workflow.completed_at - workflow.started_at).total_seconds()})
            
        except Exception as e:
            logger.error(f"Error processing workflow {workflow_id}: {str(e)}")
            self._handle_workflow_error(workflow_id, str(e))
            
    async def _process_scraping(self, workflow_id: str):
        """Scrape content from all sources."""
        step = self._create_processing_step(workflow_id, "scraping")
        
        try:
            sources = self.db.query(RAGSource).filter_by(workflow_id=workflow_id).all()
            total_sources = len(sources)
            
            for idx, source in enumerate(sources):
                await self._scrape_source(source)
                
                # Update progress
                progress = ((idx + 1) / total_sources) * 20  # Scraping is 20% of total
                self._update_workflow_progress(workflow_id, progress)
                
            self._complete_processing_step(step)
            
        except Exception as e:
            self._fail_processing_step(step, str(e))
            raise
            
    async def _scrape_source(self, source: RAGSource):
        """Scrape content from a single source."""
        source.started_at = datetime.utcnow()
        source.status = SourceStatus.PROCESSING
        self.db.commit()
        
        try:
            if source.source_type == SourceType.URL:
                content, metadata = await self._scrape_url(source.source)
            elif source.source_type == SourceType.YOUTUBE_VIDEO:
                content, metadata = await self._scrape_youtube_video(source.source)
            elif source.source_type == SourceType.YOUTUBE_CHANNEL:
                content, metadata = await self._scrape_youtube_channel(source.source)
            elif source.source_type == SourceType.DOCUMENT:
                content, metadata = await self._process_document(source.source)
            elif source.source_type == SourceType.EXTERNAL_SEARCH:
                content, metadata = await self._scrape_external_search(source.source)
            else:
                raise ValueError(f"Unknown source type: {source.source_type}")
                
            # Store scraped content temporarily in metadata
            source.metadata = {**source.metadata, **metadata, "raw_content": content}
            source.status = SourceStatus.COMPLETED
            source.completed_at = datetime.utcnow()
            
        except Exception as e:
            source.status = SourceStatus.FAILED
            source.error_message = str(e)
            source.retry_count += 1
            
        self.db.commit()
        
    async def _scrape_url(self, url: str) -> Tuple[str, Dict[str, Any]]:
        """Scrape content from a URL using Jina Reader."""
        try:
            # Try Jina Reader first for better content extraction
            result = await self.jina_service.extract_content_from_url(
                url=url,
                return_format="text"
            )
            
            if result["success"] and result["content"]:
                metadata = {
                    "url": url,
                    "extracted_at": result["extracted_at"],
                    "extraction_method": "jina_reader",
                }
                return result["content"], metadata
            else:
                logger.warning(f"Jina Reader failed for {url}, falling back to BeautifulSoup")
                
        except Exception as e:
            logger.error(f"Error with Jina Reader for {url}: {e}")
        
        # Fallback to BeautifulSoup
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                html = await response.text()
                
        soup = BeautifulSoup(html, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
            
        # Get text
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        metadata = {
            "title": soup.title.string if soup.title else url,
            "url": url,
            "scraped_at": datetime.utcnow().isoformat(),
            "extraction_method": "beautifulsoup",
        }
        
        return text, metadata
        
    async def _scrape_youtube_video(self, url: str) -> Tuple[str, Dict[str, Any]]:
        """Scrape YouTube video transcript using Jina."""
        try:
            # Use Jina to extract YouTube transcript
            result = await self.jina_service.extract_youtube_transcript(
                video_url=url,
                language="en"
            )
            
            if result["success"] and result["transcript"]:
                metadata = {
                    "url": url,
                    **result.get("metadata", {}),
                    "extraction_method": "jina_youtube",
                }
                return result["transcript"], metadata
            else:
                logger.warning(f"Jina YouTube extraction failed for {url}")
                
        except Exception as e:
            logger.error(f"Error extracting YouTube transcript with Jina: {e}")
        
        # Fallback: extract video ID for future processing
        parsed_url = urlparse(url)
        if parsed_url.hostname in ['www.youtube.com', 'youtube.com']:
            video_id = parse_qs(parsed_url.query).get('v', [None])[0]
        elif parsed_url.hostname == 'youtu.be':
            video_id = parsed_url.path[1:]
        else:
            raise ValueError(f"Invalid YouTube URL: {url}")
            
        # TODO: Implement fallback with yt-dlp or YouTube API
        return f"YouTube video transcript for {video_id} (extraction failed)", {
            "video_id": video_id,
            "url": url,
            "title": f"Video {video_id}",
            "duration": "unknown",
            "extraction_method": "fallback",
        }
        
    async def _scrape_youtube_channel(self, url: str) -> Tuple[str, Dict[str, Any]]:
        """Scrape YouTube channel videos."""
        # TODO: Implement YouTube channel scraping
        # This would involve getting list of videos and processing each
        return f"YouTube channel content from {url}", {
            "channel_url": url,
            "video_count": 0,
        }
        
    async def _scrape_external_search(self, search_query: str) -> Tuple[str, Dict[str, Any]]:
        """Scrape content using external search with keywords."""
        try:
            # Parse the search query - expecting either keywords or a JSON string
            if search_query.startswith("{") or search_query.startswith("["):
                # JSON format with keywords and context
                import json
                query_data = json.loads(search_query)
                keywords = query_data.get("keywords", [])
                context = query_data.get("context", None)
                max_results = query_data.get("max_results", 10)
            else:
                # Simple comma-separated keywords
                keywords = [k.strip() for k in search_query.split(",") if k.strip()]
                context = None
                max_results = 10
            
            logger.info(f"External search starting with keywords: {keywords}")
            if context:
                logger.info(f"Search context provided: {context}")
            
            # Calculate query distribution
            max_queries = min(3, max(1, max_results // 3))
            max_results_per_query = max(3, max_results // max_queries)
            
            logger.info(f"Query distribution: {max_queries} queries, {max_results_per_query} results per query")
            
            # Use the external search service to generate formatted queries and search
            results = await self.external_search_service.search_with_formatted_queries(
                keywords=keywords,
                context=context,
                max_queries=max_queries,
                max_results_per_query=max_results_per_query
            )
            
            logger.info(f"External search returned {len(results)} results")
            
            # Process results to extract content
            processed_results = await self.external_search_service.extract_content_from_results(results)
            
            logger.info(f"Processed {len(processed_results)} results with extracted content")
            
            # Combine all content
            content_parts = []
            sources_metadata = []
            formatted_queries = set()  # Track unique queries used
            
            for i, result in enumerate(processed_results):
                content_parts.append(f"=== Result {i+1}: {result['title']} ===\n{result['content']}")
                
                # Extract the query used for this result
                if result.get("metadata", {}).get("search_query"):
                    formatted_queries.add(result["metadata"]["search_query"])
                
                sources_metadata.append({
                    "url": result["url"],
                    "title": result["title"],
                    "search_query": result.get("metadata", {}).get("search_query", ""),
                    "metadata": result.get("metadata", {})
                })
            
            combined_content = "\n\n".join(content_parts)
            
            # Log the transformation from keywords to queries
            logger.info(f"Keywords transformed into {len(formatted_queries)} unique queries:")
            for query in formatted_queries:
                logger.info(f"  - Query: '{query}'")
            
            metadata = {
                "search_keywords": keywords,
                "search_context": context,
                "formatted_queries": list(formatted_queries),
                "results_count": len(processed_results),
                "sources": sources_metadata,
                "scraped_at": datetime.utcnow().isoformat(),
            }
            
            logger.info(f"External search completed successfully with {len(processed_results)} results")
            
            return combined_content, metadata
            
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON format in search query: {str(e)}")
            raise ValueError(f"Invalid search query format. Expected JSON or comma-separated keywords: {str(e)}")
        except Exception as e:
            logger.error(f"Error in external search: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to perform external search: {str(e)}")
        
    async def _process_document(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process a document file."""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._process_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self._process_docx(file_path)
        elif file_extension == '.txt':
            return self._process_txt(file_path)
        elif file_extension == '.csv':
            return self._process_csv(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
            
    def _process_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF."""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        metadata = {
            "file_name": os.path.basename(file_path),
            "file_type": "pdf",
            "num_pages": num_pages,
        }
        
        return text, metadata
        
    def _process_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX."""
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        metadata = {
            "file_name": os.path.basename(file_path),
            "file_type": "docx",
            "num_paragraphs": len(doc.paragraphs),
        }
        
        return text, metadata
        
    def _process_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
            
        metadata = {
            "file_name": os.path.basename(file_path),
            "file_type": "txt",
            "file_size": os.path.getsize(file_path),
        }
        
        return text, metadata
        
    def _process_csv(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from CSV."""
        text_parts = []
        with open(file_path, 'r', encoding='utf-8') as file:
            csv_reader = csv.reader(file)
            for row in csv_reader:
                text_parts.append(" | ".join(row))
                
        text = "\n".join(text_parts)
        
        metadata = {
            "file_name": os.path.basename(file_path),
            "file_type": "csv",
            "num_rows": len(text_parts),
        }
        
        return text, metadata
        
    async def _process_chunking(self, workflow_id: str):
        """Chunk all source content."""
        step = self._create_processing_step(workflow_id, "chunking")
        
        try:
            workflow = self.db.query(RAGWorkflow).filter_by(id=workflow_id).first()
            sources = self.db.query(RAGSource).filter_by(
                workflow_id=workflow_id,
                status=SourceStatus.COMPLETED
            ).all()
            
            chunk_size = workflow.parameters.get("chunkSize", 512)
            overlap = workflow.parameters.get("overlap", 50)
            
            total_chunks = 0
            for source in sources:
                content = source.metadata.get("raw_content", "")
                chunks = self._create_chunks(content, chunk_size, overlap)
                
                for idx, chunk_text in enumerate(chunks):
                    chunk = RAGChunk(
                        source_id=source.id,
                        chunk_index=idx,
                        content=chunk_text,
                        metadata={
                            "source_title": source.metadata.get("title", ""),
                            "chunk_size": chunk_size,
                            "overlap": overlap,
                        },
                        token_count=self._count_tokens(chunk_text)
                    )
                    self.db.add(chunk)
                    total_chunks += 1
                    
                # Update source stats
                source.stats = {
                    **source.stats,
                    "chunks": len(chunks),
                    "total_tokens": sum(self._count_tokens(c) for c in chunks),
                }
                
            self.db.commit()
            
            # Update workflow stats
            workflow.stats = {
                **workflow.stats,
                "total_chunks": total_chunks,
            }
            self.db.commit()
            
            self._complete_processing_step(step)
            
        except Exception as e:
            self._fail_processing_step(step, str(e))
            raise
            
    def _create_chunks(self, text: str, chunk_size: int, overlap: int) -> List[str]:
        """Create overlapping chunks from text."""
        # Simple word-based chunking
        words = text.split()
        chunks = []
        
        if not words:
            return chunks
            
        stride = max(1, chunk_size - overlap)
        
        for i in range(0, len(words), stride):
            chunk_words = words[i:i + chunk_size]
            chunk = " ".join(chunk_words)
            chunks.append(chunk)
            
            # Stop if we've reached the end
            if i + chunk_size >= len(words):
                break
                
        return chunks
        
    def _count_tokens(self, text: str) -> int:
        """Count tokens in text using tiktoken."""
        try:
            encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")
            return len(encoding.encode(text))
        except Exception:
            # Fallback to word count estimation
            return len(text.split())
            
    async def _process_embedding(self, workflow_id: str):
        """Generate embeddings for all chunks."""
        step = self._create_processing_step(workflow_id, "embedding")
        
        try:
            workflow = self.db.query(RAGWorkflow).filter_by(id=workflow_id).first()
            embedding_model = workflow.parameters.get("embeddingModel", EmbeddingModel.ADA_002)
            
            # Get all chunks
            chunks = self.db.query(RAGChunk).join(RAGSource).filter(
                RAGSource.workflow_id == workflow_id
            ).all()
            
            total_chunks = len(chunks)
            embeddings_created = 0
            
            # Process in batches
            batch_size = 100
            for i in range(0, total_chunks, batch_size):
                batch = chunks[i:i + batch_size]
                batch_texts = [chunk.content for chunk in batch]
                
                # Generate embeddings
                embeddings = await self._generate_embeddings(batch_texts, embedding_model)
                
                # Store embeddings
                for chunk, embedding in zip(batch, embeddings):
                    embedding_record = RAGEmbedding(
                        workflow_id=workflow_id,
                        chunk_id=chunk.id,
                        vector=self._serialize_vector(embedding),
                        vector_dimension=len(embedding),
                        embedding_model=embedding_model,
                        metadata={
                            "chunk_index": chunk.chunk_index,
                            "source_id": str(chunk.source_id),
                        }
                    )
                    self.db.add(embedding_record)
                    embeddings_created += 1
                    
                self.db.commit()
                
                # Update progress
                progress = 20 + ((embeddings_created / total_chunks) * 40)  # Embedding is 40% of total
                self._update_workflow_progress(workflow_id, progress)
                
            # Update workflow stats
            workflow.stats = {
                **workflow.stats,
                "embeddings": embeddings_created,
            }
            self.db.commit()
            
            self._complete_processing_step(step)
            
        except Exception as e:
            self._fail_processing_step(step, str(e))
            raise
            
    async def _generate_embeddings(self, texts: List[str], model: str) -> List[List[float]]:
        """Generate embeddings using Jina AI or OpenAI as fallback."""
        # Check if model is a Jina model
        if model.startswith("jina-"):
            try:
                # Use Jina AI for embeddings
                embeddings = await self.jina_service.generate_embeddings(
                    texts=texts,
                    model=model,
                    dimensions=1024,
                    normalize=True,
                    batch_size=32
                )
                
                # Calculate and log cost
                total_tokens = sum(len(self.tokenizer.encode(text)) for text in texts)
                estimated_cost = (total_tokens / 1_000_000) * 0.02  # $0.02 per 1M tokens
                logger.info(f"Generated {len(embeddings)} embeddings using Jina AI. Estimated cost: ${estimated_cost:.4f}")
                
                return embeddings
                
            except Exception as e:
                logger.error(f"Error generating embeddings with Jina: {e}")
                logger.info("Falling back to OpenAI embeddings")
                model = "text-embedding-ada-002"  # Fallback model
        
        # Use OpenAI for embeddings (fallback or if not Jina model)
        response = await self.openai_client.embeddings.create(
            input=texts,
            model=model
        )
        
        embeddings = [item.embedding for item in response.data]
        
        # Update cost tracking for OpenAI
        # OpenAI pricing: $0.0001 per 1K tokens for ada-002
        total_tokens = sum(len(self.tokenizer.encode(text)) for text in texts)
        estimated_cost = (total_tokens / 1000) * 0.0001
        logger.info(f"Generated {len(embeddings)} embeddings using OpenAI. Estimated cost: ${estimated_cost:.4f}")
        
        return embeddings
        
    def _serialize_vector(self, vector: List[float]) -> bytes:
        """Serialize vector to binary format."""
        return np.array(vector, dtype=np.float32).tobytes()
        
    def _deserialize_vector(self, vector_bytes: bytes) -> List[float]:
        """Deserialize vector from binary format."""
        return np.frombuffer(vector_bytes, dtype=np.float32).tolist()
        
    async def _process_indexing(self, workflow_id: str):
        """Index embeddings in vector store."""
        step = self._create_processing_step(workflow_id, "indexing")
        
        try:
            workflow = self.db.query(RAGWorkflow).filter_by(id=workflow_id).first()
            vector_store = workflow.parameters.get("vectorStore", VectorStore.PINECONE)
            
            # Get all embeddings
            embeddings = self.db.query(RAGEmbedding).filter_by(workflow_id=workflow_id).all()
            
            # Index based on vector store type
            if vector_store == VectorStore.PINECONE:
                await self._index_to_pinecone(workflow_id, embeddings)
            elif vector_store == VectorStore.CHROMA:
                await self._index_to_chroma(workflow_id, embeddings)
            elif vector_store == VectorStore.WEAVIATE:
                await self._index_to_weaviate(workflow_id, embeddings)
            elif vector_store == VectorStore.QDRANT:
                await self._index_to_qdrant(workflow_id, embeddings)
            elif vector_store == VectorStore.FAISS:
                await self._index_to_faiss(workflow_id, embeddings)
            else:
                raise ValueError(f"Unknown vector store: {vector_store}")
                
            # Update progress
            self._update_workflow_progress(workflow_id, 80)  # Indexing complete at 80%
            
            # Calculate index size (placeholder)
            index_size_mb = len(embeddings) * 0.1  # Rough estimate
            workflow.stats = {
                **workflow.stats,
                "indexSize": f"{index_size_mb:.1f} MB",
            }
            self.db.commit()
            
            self._complete_processing_step(step)
            
        except Exception as e:
            self._fail_processing_step(step, str(e))
            raise
            
    async def _index_to_pinecone(self, workflow_id: str, embeddings: List[RAGEmbedding]):
        """Index embeddings to Pinecone."""
        # TODO: Implement Pinecone integration
        logger.info(f"Indexing {len(embeddings)} embeddings to Pinecone for workflow {workflow_id}")
        
    async def _index_to_chroma(self, workflow_id: str, embeddings: List[RAGEmbedding]):
        """Index embeddings to ChromaDB."""
        # TODO: Implement ChromaDB integration
        logger.info(f"Indexing {len(embeddings)} embeddings to ChromaDB for workflow {workflow_id}")
        
    async def _index_to_weaviate(self, workflow_id: str, embeddings: List[RAGEmbedding]):
        """Index embeddings to Weaviate."""
        # TODO: Implement Weaviate integration
        logger.info(f"Indexing {len(embeddings)} embeddings to Weaviate for workflow {workflow_id}")
        
    async def _index_to_qdrant(self, workflow_id: str, embeddings: List[RAGEmbedding]):
        """Index embeddings to Qdrant."""
        # TODO: Implement Qdrant integration
        logger.info(f"Indexing {len(embeddings)} embeddings to Qdrant for workflow {workflow_id}")
        
    async def _index_to_faiss(self, workflow_id: str, embeddings: List[RAGEmbedding]):
        """Index embeddings to FAISS."""
        # TODO: Implement FAISS integration
        logger.info(f"Indexing {len(embeddings)} embeddings to FAISS for workflow {workflow_id}")
        
    async def _process_validation(self, workflow_id: str):
        """Validate the indexed data."""
        step = self._create_processing_step(workflow_id, "validating")
        
        try:
            # Perform validation checks
            workflow = self.db.query(RAGWorkflow).filter_by(id=workflow_id).first()
            
            # Check that all sources were processed
            sources = self.db.query(RAGSource).filter_by(workflow_id=workflow_id).all()
            failed_sources = [s for s in sources if s.status == SourceStatus.FAILED]
            
            if failed_sources:
                logger.warning(f"Workflow {workflow_id} has {len(failed_sources)} failed sources")
                
            # Verify embedding count
            embedding_count = self.db.query(RAGEmbedding).filter_by(workflow_id=workflow_id).count()
            
            # Update final stats
            workflow.stats = {
                **workflow.stats,
                "validation": {
                    "total_sources": len(sources),
                    "failed_sources": len(failed_sources),
                    "total_embeddings": embedding_count,
                    "validated_at": datetime.utcnow().isoformat(),
                }
            }
            
            # Update progress
            self._update_workflow_progress(workflow_id, 100)
            
            self.db.commit()
            self._complete_processing_step(step)
            
        except Exception as e:
            self._fail_processing_step(step, str(e))
            raise
            
    def _create_processing_step(self, workflow_id: str, step_name: str) -> RAGProcessingStep:
        """Create a new processing step record."""
        step = RAGProcessingStep(
            workflow_id=workflow_id,
            step_name=step_name,
            status="processing",
            started_at=datetime.utcnow()
        )
        self.db.add(step)
        self.db.commit()
        return step
        
    def _complete_processing_step(self, step: RAGProcessingStep):
        """Mark a processing step as completed."""
        step.status = "completed"
        step.completed_at = datetime.utcnow()
        step.duration_seconds = (step.completed_at - step.started_at).total_seconds()
        step.progress = 100.0
        self.db.commit()
        
    def _fail_processing_step(self, step: RAGProcessingStep, error_message: str):
        """Mark a processing step as failed."""
        step.status = "failed"
        step.error_message = error_message
        step.completed_at = datetime.utcnow()
        step.duration_seconds = (step.completed_at - step.started_at).total_seconds()
        self.db.commit()
        
    def _update_workflow_progress(self, workflow_id: str, progress: float):
        """Update workflow progress."""
        self.db.execute(
            update(RAGWorkflow)
            .where(RAGWorkflow.id == workflow_id)
            .values(progress=progress, updated_at=datetime.utcnow())
        )
        self.db.commit()
        
    def _log_event(self, workflow_id: str, event_type: str, event_data: Dict[str, Any]):
        """Log a workflow event."""
        event = RAGWorkflowEvent(
            workflow_id=workflow_id,
            event_type=event_type,
            event_data=event_data
        )
        self.db.add(event)
        self.db.commit()
        
    def _handle_workflow_error(self, workflow_id: str, error_message: str):
        """Handle workflow error."""
        workflow = self.db.query(RAGWorkflow).filter_by(id=workflow_id).first()
        if workflow:
            workflow.status = WorkflowStatus.FAILED
            self.db.commit()
            
        self._log_event(workflow_id, "failed", {"error": error_message})


class YouTubeTranscriptService:
    """Service for extracting YouTube transcripts."""
    
    def __init__(self):
        # TODO: Initialize YouTube API client
        pass
        
    async def get_video_transcript(self, video_id: str) -> Tuple[str, Dict[str, Any]]:
        """Get transcript for a YouTube video."""
        # TODO: Implement YouTube transcript extraction
        # This would use youtube-transcript-api or similar
        return f"Transcript for video {video_id}", {
            "video_id": video_id,
            "duration": "unknown",
            "language": "en",
        }
        
    async def get_channel_videos(self, channel_url: str) -> List[str]:
        """Get list of video URLs from a YouTube channel."""
        # TODO: Implement YouTube channel video listing
        return []